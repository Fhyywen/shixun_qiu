import numpy as np
from typing import List, Dict, Any
import os
import pandas as pd
from sentence_transformers import SentenceTransformer
from config import Config  # 正确的导入方式


class DataProcessor:
    def __init__(self, model_name: str = None):
        if model_name is None:
            config = Config()
            self.model_name = config.EMBEDDING_MODEL
        else:
            self.model_name = model_name

        self.model = None
        print(f"初始化数据处理器，使用模型: {self.model_name}")

    def _load_model(self):
        """延迟加载模型"""
        if self.model is None:
            print("正在加载嵌入模型...")
            self.model = SentenceTransformer(self.model_name)
            print("模型加载完成")

    def check_existing_documents(self, vector_store, documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        检查并向量库中已存在的文档，避免重复处理
        """
        unique_documents = []
        for doc in documents:
            # 基于文件路径或其他唯一标识符检查是否已存在
            try:
                existing = vector_store.get(
                    where={"source": doc["source"]}
                    # 移除 include=["documents"]，因为默认会返回 ids
                    # 如果确实需要 documents 内容，可以使用 include=["documents", "metadatas"]
                )

                if len(existing['ids']) == 0:
                    unique_documents.append(doc)
                else:
                    print(f"文档已存在，跳过: {doc['source']}")
            except Exception as e:
                print(f"检查文档存在性时出错: {e}")
                # 出错时仍然添加文档，避免遗漏
                unique_documents.append(doc)

        return unique_documents

    def load_documents(self, data_path: str) -> List[Dict[str, Any]]:
        """加载多种格式的文档"""
        documents = []

        if not os.path.exists(data_path):
            print(f"数据路径不存在: {data_path}")
            return documents

        for root, _, files in os.walk(data_path):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    if file.endswith(('.txt', '.md', '.rst')):
                        content = self._load_text_file(file_path)
                    elif file.endswith('.csv'):
                        content = self._load_csv_file(file_path)
                    elif file.endswith(('.xlsx', '.xls')):
                        content = self._load_excel_file(file_path)
                    elif file.endswith('.docx'):  # 添加Word文档支持
                        content = self._load_word_file(file_path)
                    elif file.endswith('.doc'):
                        content = self._load_doc_file(file_path)
                    elif file.endswith('.pdf'):
                        content = self._load_pdf_file(file_path)
                    else:
                        print(f"跳过不支持的文件格式: {file}")
                        continue

                    documents.append({
                        "content": content,
                        "source": file_path,
                        "type": os.path.splitext(file)[1]
                    })
                    print(f"成功加载文件: {file}")

                except Exception as e:
                    print(f"加载文件 {file_path} 时出错: {e}")

        print(f"共加载 {len(documents)} 个文档")
        return documents

    def _load_doc_file(self, file_path: str) -> str:
        """加载.doc文件并转换为文本"""
        try:
            # 尝试使用win32com读取.doc文件
            import win32com.client

            text = f"Word文档(.doc): {os.path.basename(file_path)}\n"
            text += "=" * 50 + "\n"

            # 创建Word应用程序对象
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False  # 不显示Word界面

            try:
                # 打开文档
                doc = word.Documents.Open(file_path)

                # 获取文档内容
                content = doc.Content.Text
                text += content

                # 关闭文档
                doc.Close()

            finally:
                # 退出Word应用程序
                word.Quit()

            return text

        except ImportError:
            # 如果win32com不可用，尝试使用antiword（需要安装antiword工具）
            try:
                import subprocess

                # 使用antiword提取文本内容
                result = subprocess.run(['antiword', file_path],
                                        capture_output=True, text=True, timeout=30)

                if result.returncode == 0:
                    text = f"Word文档(.doc): {os.path.basename(file_path)}\n"
                    text += "=" * 50 + "\n"
                    text += result.stdout
                    return text
                else:
                    return f"使用antiword读取.doc文件失败: {result.stderr}"

            except (FileNotFoundError, subprocess.TimeoutExpired):
                return f"错误: 无法读取.doc文件 {file_path}。请确保已安装win32com或antiword"

        except Exception as e:
            return f"读取.doc文件 {file_path} 出错: {str(e)}"

    def _load_text_file(self, file_path: str) -> str:
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _load_csv_file(self, file_path: str) -> str:
        """加载CSV文件并转换为文本"""
        try:
            df = pd.read_csv(file_path)
            # 将DataFrame转换为描述性文本
            text = f"CSV文件: {os.path.basename(file_path)}\n"
            text += f"列名: {', '.join(df.columns)}\n"
            text += f"行数: {len(df)}\n"
            text += f"前3行数据:\n{df.head(3).to_string()}\n"
            return text
        except Exception as e:
            return f"读取CSV文件出错: {str(e)}"

    def _load_word_file(self, file_path: str) -> str:
        """加载Word文档文件并转换为文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = f"Word文档: {os.path.basename(file_path)}\n"
            text += "=" * 50 + "\n"

            # 提取段落文本
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text += "\n".join(paragraphs)

            # 可选：提取表格内容
            if doc.tables:
                text += "\n\n文档中的表格:\n"
                for i, table in enumerate(doc.tables):
                    text += f"\n表格 {i + 1}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        text += row_text + "\n"

            return text
        except ImportError:
            return f"错误: 未安装python-docx库，无法读取Word文档 {file_path}"
        except Exception as e:
            return f"读取Word文档 {file_path} 出错: {str(e)}"

    def _load_excel_file(self, file_path: str) -> str:
        """加载Excel文件并转换为文本"""
        try:
            df = pd.read_excel(file_path)
            text = f"Excel文件: {os.path.basename(file_path)}\n"
            text += f"列名: {', '.join(df.columns)}\n"
            text += f"行数: {len(df)}\n"
            text += f"前3行数据:\n{df.head(3).to_string()}\n"
            return text
        except Exception as e:
            return f"读取Excel文件出错: {str(e)}"

    def chunk_documents(self, documents: List[Dict[str, Any]], chunk_size: int = 500, chunk_overlap: int = 50) -> List[
        Dict[str, Any]]:
        """将文档分块"""
        chunks = []

        for doc in documents:
            content = doc["content"]
            source = doc["source"]

            # 简单的文本分块
            words = content.split()
            for i in range(0, len(words), chunk_size - chunk_overlap):
                chunk_words = words[i:i + chunk_size]
                chunk_text = " ".join(chunk_words)

                chunks.append({
                    "content": chunk_text,
                    "source": source,
                    "chunk_index": len(chunks)
                })

        print(f"文档分块完成，共 {len(chunks)} 个块")
        return chunks

    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        """生成文本嵌入向量"""
        self._load_model()  # 确保模型已加载

        print(f"正在为 {len(texts)} 个文本生成嵌入...")
        embeddings = self.model.encode(texts)
        print("嵌入生成完成")
        return embeddings

    def _load_pdf_file(self, file_path: str) -> str:
        """加载PDF文件文本。优先抽取文本型PDF；扫描件建议结合OCR（可后续增强）。"""
        try:
            import pdfplumber
        except ImportError:
            return f"错误: 未安装 pdfplumber，无法读取PDF文件 {os.path.basename(file_path)}"

        texts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ''
                    if page_text.strip():
                        texts.append(page_text)
        except Exception as e:
            return f"读取PDF文件出错: {str(e)}"

        header = f"PDF文件: {os.path.basename(file_path)}\n" + ("=" * 50) + "\n"
        body = "\n\n".join(texts).strip()
        if not body:
            # 未抽取到文本（可能是扫描件）- 尝试OCR回退
            ocr_text = self._ocr_pdf(file_path)
            if ocr_text:
                return header + ocr_text
            return header + "(未从PDF中提取到文本，且OCR回退失败。请检查PaddleOCR/显卡驱动/依赖安装。)"
        return header + body

    def _ensure_paddle_ocr(self):
        """确保 PaddleOCR 可用（按需安装）。"""
        flag_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), '.paddle_installed')
        try:
            # 已可导入则直接返回
            try:
                from paddleocr import PaddleOCR  # noqa: F401
                return True
            except Exception:
                pass

            # 有安装标记但仍不可用，继续尝试导入失败则返回False
            if os.path.exists(flag_path):
                try:
                    from paddleocr import PaddleOCR  # noqa: F401
                    return True
                except Exception:
                    return False

            import subprocess, sys
            python_exe = sys.executable
            # 优先安装 GPU 版本（与 app 中保持一致），失败则不抛出
            subprocess.run([
                python_exe, '-m', 'pip', 'install', 'paddlepaddle-gpu==3.2.0', '-i',
                'https://www.paddlepaddle.org.cn/packages/stable/cu118/'
            ], check=False)
            subprocess.run([python_exe, '-m', 'pip', 'install', 'paddleocr[all]'], check=False)

            # 再次尝试导入
            try:
                from paddleocr import PaddleOCR  # noqa: F401
                with open(flag_path, 'w', encoding='utf-8') as f:
                    f.write('installed')
                return True
            except Exception:
                return False
        except Exception:
            return False

    def _ocr_pdf(self, file_path: str) -> str:
        """将PDF渲染为图像并用 PaddleOCR 识别，返回拼接文本。"""
        # 确保 OCR 依赖就绪
        if not self._ensure_paddle_ocr():
            return ''

        try:
            import fitz  # PyMuPDF
            from paddleocr import PaddleOCR
        except Exception:
            return ''

        try:
            ocr = PaddleOCR(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
            )

            doc = fitz.open(file_path)
            page_texts = []
            for page in doc:
                # 渲染为图像（缩放以提升清晰度）
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_bytes = pix.tobytes('png')

                # 将字节喂给 OCR（PaddleOCR 支持 numpy 数组/路径；这里用临时字节转换）
                import numpy as np
                import cv2
                nparr = np.frombuffer(img_bytes, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                if img is None:
                    continue

                result = ocr.ocr(img, cls=False)
                if not result:
                    continue
                # result 结构: [ [ [box, (text, score)], ... ] ]
                lines = []
                for line in result[0] if isinstance(result, list) and len(result) > 0 else []:
                    try:
                        text = line[1][0]
                        if text:
                            lines.append(text)
                    except Exception:
                        continue
                if lines:
                    page_texts.append("\n".join(lines))

            return "\n\n".join(page_texts).strip()
        except Exception:
            return ''